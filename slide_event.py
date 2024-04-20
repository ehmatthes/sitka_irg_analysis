"""Model for landslide events."""

import datetime, json

import pytz

from docx import Document
from docx.shared import Inches


class SlideEvent:

    def __init__(self, dt_slide=None):

        # datetime of slide event
        self.dt_slide = dt_slide
        
        self.desc_location = ''
        self.name = ''

        self.power_outage = None
        self.fatalities = None
        
        self.gps_location = None
        
        self.urls = []


    def summarize_slide(self):
        """Summarize all information about the slide."""
        pass
        

    def __str__(self):
        return self.name


    def as_html(self):
        """Return an html snippet that summarizes all information about
        a slide event.
        Can be used to generate a web page summarizing all events.

        # DEV: This should probably be @staticmethod.
        """
        html_str = '<div>'
        html_str += f"\n  <h3>{self.name}</h3>"
        html_str += f"\n  <p>{self.dt_slide.strftime('%m/%d/%Y %H:%M:%S')}</p>"
        html_str += f"\n  <p>Location: {self.desc_location}</p>"
        html_str += f"\n  <p>About this slide:</p>"
        html_str += f"\n  <ul>"
        for url in self.urls:
            html_str += f"\n    <li><a href='{url}'>{url}</a></li>"
        html_str += f"\n  </ul>"
        html_str += f"\n</div>"

        return html_str


    @classmethod
    def load_slides(se_cls, data_file):
        """Load slides from json string format into SlideEvent objects.
        Return list of SlideEvent objects.
        """
        with open(data_file) as f:
            slides_json = json.load(f)

        # Convert json to class instances.
        slide_objects = []
        for slide_dict in slides_json:
            new_slide = se_cls()
            new_slide.__dict__ = slide_dict
            new_slide.dt_slide = datetime.datetime.strptime(
                    new_slide.dt_slide, '%Y-%m-%d %H:%M:%S+00:00')
            new_slide.dt_slide = new_slide.dt_slide.replace(tzinfo=pytz.utc)
            slide_objects.append(new_slide)

        return slide_objects


    @staticmethod
    def get_web_page(known_slides):
        """Given a list of known slide instances, generate a web page
        summarizing all slides."""
        html_str = "<html><body>"
        html_str += "\n<div><h1>Landslides on the Sitka Road System</h1></div>"
        html_str += "\n\n"

        for slide in known_slides:
            html_str += slide.as_html()

        html_str += "\n\n</body></html>"

        return html_str


    @staticmethod
    def generate_word_doc(known_slides, filename):
        """Given a list of known slide instances, generate a word doc
        summarizing all slides."""
        document = Document()

        from my_docx_utils import add_hyperlink

        for slide in known_slides:
            run = document.add_paragraph(style='Heading 2').add_run(slide.name)
            font = run.font
            font.name = 'Calibri Light'

            p = document.add_paragraph()
            p.paragraph_format.space_after = 6
            run = p.add_run(slide.dt_slide.strftime('%m/%d/%Y %H:%M:%S'))
            run.font.name = 'Garamond'

            p = document.add_paragraph()
            p.paragraph_format.space_after = 6
            run = p.add_run(slide.desc_location)
            run.font.name = 'Garamond'

            for url in slide.urls:
                # run = document.add_paragraph().add_run(url)
                # run.font.name = 'Garamond'
                p = document.add_paragraph()
                p.style = 'List Bullet'
                run = p.add_run()
                run.font.name = 'Garamond'
                add_hyperlink(p, url, url, font_name='Garamond',
                    underline_hyperlink=False, indent=0.5, color=False)

        document.save(filename)


if __name__ == '__main__':

    known_slides = []

    # Create the known events here, for now.

    """Further research

    November 2005 - (City Shops Slide)
    In November, 2005, Sitka experienced two catastrophic landslide events
    which closed Halibut Point Road (HPR), owned by the State of Alaska.
    HPR is Sitka's major artery from the ferry terminal and largest
    residential section of the community into downtown Sitka. This slide area,
    located between Davidoff Street and the Old City Shops, has experienced
    several landslides over many years. The slides involved several private
    residences, municipal road rights- of-way and the State Halibut Point Road
    facility. In addition to closing the road and damaging various private
    residences, the former State maintenance building was reduced to rubble,
    and the State highway was buried under tons of debris.
    URL: http://www.cityofsitka.com/government/departments/planning/documents/CompPlanNovember06.pdf
    
    Cascade Creek Road slide that moved part of house
    """

    # Beaver Lake Slide 11/2011 (wind and snowmelt?)
    #  Eyewitnesses may be able to provide a time for this slide.
    #  May want to look at snowmelt and wind factors for this slide.
    new_slide = SlideEvent()
    # 11/12/2011 10:00:00 AKST; should be 19:00:00 UTC
    new_slide.dt_slide = datetime.datetime(2011, 11, 12, 19, 0, 0, tzinfo=pytz.utc)
    new_slide.name = 'Beaver Lake Slide 11/2011 (wind and snowmelt?)'
    new_slide.desc_location = 'Beaver Lake, Bear Mountain shoreline'
    new_slide.fatalities = 0
    new_slide.power_outage = None
    new_slide.urls.append('https://www.kcaw.org/2011/12/19/mass-wasting-event-destroys-popular-sitka-trail/')
    known_slides.append(new_slide)

    # Redoubt Slide 5/2013 (not on Sitka road system)
    new_slide = SlideEvent()
    # 5/13/2013 11:00:00 AKDT; should be 19:00:00 UTC
    new_slide.dt_slide = datetime.datetime(2013, 5, 13, 19, 0, 0, tzinfo=pytz.utc)
    new_slide.name = 'Redoubt Slide 5/2013 (not on Sitka road system)'
    new_slide.desc_location = 'Redoubt Lake, near Redoubt Lake Cabin'
    new_slide.fatalities = 0
    new_slide.power_outage = None
    new_slide.urls.append('https://www.kcaw.org/2013/05/13/couple-escapes-as-landslide-destroys-cabin/')
    known_slides.append(new_slide)

    # Starrigavan Slide 9/2014 (time of slide unknown)
    new_slide = SlideEvent()
    # Currently using 9/18/2014 12:00:00 AKDT; this would be 20:00:00 UTC
    new_slide.dt_slide = datetime.datetime(2014, 9, 18, 20, 0, 0, tzinfo=pytz.utc)
    new_slide.name = 'Starrigavan Slide 9/2014 (time of slide unknown)'
    new_slide.desc_location = 'Starrigavan Valley'
    new_slide.fatalities = 0
    new_slide.power_outage = None
    new_slide.urls.append('https://www.kcaw.org/2014/09/24/landslide-destroys-starrigavan-restoration-projects/')
    new_slide.urls.append('http://www.sitkanature.org/wordpress/2014/09/26/starrigavan-landslide/')
    known_slides.append(new_slide)

    # South Kramer Slide 8/2015
    kramer_slide = SlideEvent()
    # 8/18/2015 9:41:00 AKDT; should be 17:41:00 UTC
    kramer_slide.dt_slide = datetime.datetime(2015, 8, 18, 17, 41, 0, tzinfo=pytz.utc)
    kramer_slide.name = 'South Kramer Slide 8/2015'
    kramer_slide.desc_location = 'South end of Kramer Ave'
    kramer_slide.fatalities = 3
    kramer_slide.urls.append('https://www.adn.com/alaska-news/article/3-missing-after-heavy-rain-prompts-landslides-sinkhole-sitka/2015/08/18/')
    kramer_slide.urls.append('https://www.kcaw.org/2015/08/18/three-landslides-prompt-sitka-to-declare-state-of-emergency/')
    kramer_slide.urls.append('https://www.cityofsitka.com/documents/Sitka_SKramerLandslideReport.pdf')
    known_slides.append(kramer_slide)

    # HPR Slide 9/2016 (minor slide)
    new_slide = SlideEvent()
    # 9/16/2016 02:20:00 AKDT; should be 10:20:00 UTC
    new_slide.dt_slide = datetime.datetime(2016, 9, 16, 10, 20, 0, tzinfo=pytz.utc)
    new_slide.name = 'HPR Slide 9/2016 (minor slide)'
    new_slide.desc_location = 'HPR, near Davidoff Street'
    new_slide.fatalities = 0
    new_slide.power_outage = None
    new_slide.urls.append('https://www.kcaw.org/2016/09/16/small-mudslide-generates-big-response-in-sitka/')
    known_slides.append(new_slide)

    # HPR Slide 9/2017
    hpr_slide = SlideEvent()
    # 9/4/2017 12:00:00 AKDT; should be 20:00:00 UTC
    hpr_slide.dt_slide = datetime.datetime(2017, 9, 4, 20, 0, 0, tzinfo=pytz.utc)
    hpr_slide.name = 'HPR Slide 9/2017'
    hpr_slide.desc_location = 'HPR, near Valhalla Drive'
    hpr_slide.fatalities = 0
    hpr_slide.power_outage = None
    hpr_slide.urls.append('https://www.kcaw.org/2017/09/04/landslide-closes-halibut-point-road-sitka/')
    hpr_slide.urls.append('https://www.kcaw.org/2017/09/04/no-injuries-sitkas-pretty-impressive-labor-day-landslide/')
    known_slides.append(hpr_slide)

    # Medvejie Slide 9/2019
    medv_slide = SlideEvent()
    # 9/20/2019 12:50:00 AKDT; should be 20:50:00 UTC
    medv_slide.dt_slide = datetime.datetime(2019, 9, 20, 20, 50, 0, tzinfo=pytz.utc)
    medv_slide.name = 'Medvejie Slide 9/2019'
    medv_slide.desc_location = 'Medvejie Hatchery'
    medv_slide.fatalities = 0
    medv_slide.power_outage = True
    medv_slide.urls.append('https://www.kcaw.org/2019/09/20/slide-cuts-off-green-lake-road-hatchery-access/')
    known_slides.append(medv_slide)

    # Sand Dollar Drive slide 11/2/2020
    sdd_slide = SlideEvent()
    # 11/01/2020 19:25:00 AKST; should be 04:25:00 11/02/20 UTC
    #   This is well-specified, from personal communication with homeowners.
    #   "It was after 7. 7:20 my best guess."
    #   "[partner] agrees 7:20-7:30"
    #  Also note timezone changed from AKDT to AKST 0200 on 11/1/2020, about 18 hours before this slide.
    sdd_slide.dt_slide = datetime.datetime(2020, 11, 2, 4, 25, 0, tzinfo=pytz.utc)
    sdd_slide.name = 'Sand Dollar Drive Slide 11/2/2020'
    sdd_slide.desc_location = 'Sand Dollar Drive'
    sdd_slide.fatalities = 0
    sdd_slide.power_outage = False
    sdd_slide.urls.append('https://www.kcaw.org/2020/11/02/back-to-back-landslides-block-sitkas-sand-dollar-drive/')
    known_slides.append(sdd_slide)

    # Second Sand Dollar Drive slide 11/2/2020
    sdd_slide_2 = SlideEvent()
    # 11/02/2020 03:00:00 AKST; should be 12:00:00 11/02/20 UTC
    #   This is less well-specified, and was probably an additional release of the first slide?
    #   From personal communication with homeowner, and KCAW article.
    #   "The second slide happened around 3am"
    #   "The city had cleared the road after 7:30, another slide happened and blocked the road again."
    sdd_slide_2.dt_slide = datetime.datetime(2020, 11, 2, 12, 0, 0, tzinfo=pytz.utc)
    sdd_slide_2.name = 'Second Sand Dollar Drive Slide 11/2/2020'
    sdd_slide_2.desc_location = 'Sand Dollar Drive'
    sdd_slide_2.fatalities = 0
    sdd_slide_2.power_outage = False
    sdd_slide_2.urls.append('https://www.kcaw.org/2020/11/02/back-to-back-landslides-block-sitkas-sand-dollar-drive/')
    known_slides.append(sdd_slide_2)

    # Olga slides 10/26/20
    olga_slide = SlideEvent()
    # 11/01/2020 18:00:00 AKST; should be 03:00:00 11/02/20 UTC
    # 10/26/2020 12:00:00 ADDT; should be 20:00:00 10/26/20 UTC
    #   The slides probably happened 12-24 hours before this time.
    olga_slide.dt_slide = datetime.datetime(2020, 10, 26, 20, 0, 0, tzinfo=pytz.utc)
    olga_slide.name = 'Olga Strait slides 10/26/20'
    olga_slide.desc_location = 'Waterways North of Sitka'
    olga_slide.fatalities = 0
    olga_slide.power_outage = False
    olga_slide.urls.append('https://www.facebook.com/groups/sitkachatters/permalink/1816612201819511/')
    known_slides.append(olga_slide)

    # Crawfish Inlet slide 8/12/23
    ci_slide = SlideEvent()
    # Heavy rains 8/12, slide noticed 8/14 by fishing crew.
    # ~4.42 inches of rain per NWS, record for the date.
    # Unsure of time, so using 2000. Not during a critical point, but associated.
    # 8/12/23 20:00:00 AKDT; should be 04:00:00 8/13/23 UTC
    ci_slide.dt_slide = datetime.datetime(2023, 8, 13, 4, 0, 0, tzinfo=pytz.utc)
    ci_slide.name = "Crawfish Inlet slide 8/12/23"
    ci_slide.desc_location = "Crawfish Inlet, South of Sitka"
    ci_slide.fatalities = 0
    ci_slide.power_outage = False
    ci_slide.urls.append("https://www.facebook.com/AlaskanLegoMinifigure/posts/pfbid0d4ArJr11angJbfEut37LaPC8am42uek2jVMssLMK4JMCa57XYT5bt1NgvPYG3LL3l")
    ci_slide.urls.append("https://sitkascience.org/atmospheric-river-event/")
    ci_slide.urls.append("https://www.kcaw.org/2023/08/15/record-rainfall-bumped-sitkas-landslide-risk-level-to-medium-on-saturday/")
    known_slides.append(ci_slide)


    # Store this as JSON, so it can be imported into plotting code.
    #  This just stores the __dict__ elements of SlideEvent. May want to use
    #  jsonpickle, or dataclasses or namedtuple.
    #  See: https://stackoverflow.com/questions/3768895/how-to-make-a-class-json-serializable

    # Can't directly store class objects, so build a list of __dict__ from
    # known_slides.
    slides_dicts = [slide.__dict__ for slide in known_slides]
    filename = 'known_slides/known_slides.json'
    with open(filename, 'w') as f:
        json.dump(slides_dicts, f, default=str, indent=2)

    # Generate an html page listing all information about known slides.
    slides_page = SlideEvent.get_web_page(known_slides)
    filename = 'known_slides/known_slides.html'
    with open(filename, 'w') as f:
        f.writelines(slides_page)

    # Generate a word doc listing information about known slides.
    SlideEvent.generate_word_doc(known_slides, 'known_slides/known_slides.docx')